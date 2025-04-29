from retriever import Retriever
import ollama
import time as time
from tqdm import tqdm
import sys
import os
from questions import robostudio_questions, rw_rapid_questions, overview_rapid_questions, rw_system_parameters_questions, all_questions
import csv
from docx import Document
from transformers import AutoTokenizer, AutoModel
from sklearn.metrics.pairwise import cosine_similarity
import torch

class ABBot:
    """
    Args:
        retriever (Retriever): An instance of the Retriever class for document retrieval.
        llm (str): The name of the LLM to use from local Ollama models.
    """

    def __init__(self, retriever, llm, filter):
        self.retriever = retriever
        self.llm = llm
        self.filter = filter
        self.chat_history = [
            {"role": "system", "content": "You are an AI assistant that provides answers based on the given context."}
        ]
         
    def _chat_(self, query):
            if self.filter == True:
                context = self._filter_(query)
                #print(f"Filtered Context: {context}")
            else:
                context = self.retriever.retrieve(query, top_k=3)
                #print(f"Unfiltered Context: {context}")

            user_message = f"Context: {context}\n\nQuestion: {query}"
            self.chat_history.append({"role": "user", "content": user_message})

            response = ollama.chat(self.llm, self.chat_history)
            answer = response["message"]["content"]

            self.chat_history.append({"role": "assistant", "content": answer})
            print(f"\nABBot: {answer}\n")


    def _benchmark_(self):
            total_time = 0
            iteration_results = []
            question_list_types = [all_questions, robostudio_questions, rw_rapid_questions, overview_rapid_questions, rw_system_parameters_questions]

            choice = input("""Enter a number 1-5 to choose a question set:
                                \n1. All Questions\n2. Robostudio Questions\n3. RW RAPID Questions\n4. Overview RAPID Questions\n5. RW System Parameters Questions\nYour choice: """)
            
            questions = question_list_types[int(choice) - 1]
            selected_list = question_list_types[int(choice) - 1]
            list_name = [name for name, value in globals().items() if value is selected_list][0]
            iterations = int(input("Enter the number of iterations to run benchmark on: "))

            print(f"\n\nRunning benchmark for {iterations} iterations on question set: {list_name}")
            print("-" * 100)
            for question in tqdm(questions):
                for i in range(0, iterations):
                    print("*" * 100)
                    print(f"Question: {question}")
                    self._reset_chat_history_()

                    start_time = time.time()

                    self._chat_(question)

                    iteration_time = time.time() - start_time

                    total_time += iteration_time

                    iteration_results.append(iteration_time)
                    print(f"Question: {question}\nIteration {i+1}: {iteration_results[i]:.2f} seconds\n\n")
                    
                iteration_results.clear()
                print("-" * 100)

            questions_answered = len(questions) * int(iterations)
            print(f"Benchmark completed -- Retriever: {self.retriever.method} used with LLM: {self.llm}\n\n")
            print(f"Average response time: {total_time / questions_answered:.2f} seconds.")
            print(f"Total execution time on {questions_answered} questions: {total_time:.2f} seconds.")
            print(f"Question set used: {list_name}")
            print(f"Number of iterations: {iterations}")

    def write_to_docx(self, llm_name, retriever_name, filter_name):
        llm_name = llm_name.replace(":", "-")
        llm_name = llm_name.replace(".", "-")
        file_path = f"answers/Answers_{llm_name}_{retriever_name}_{filter_name}.docx"
        print(file_path)
        if not os.path.exists(file_path):
            doc = Document()
            doc.save(file_path)
        else:
            doc = Document(file_path)
        for question in tqdm(all_questions):
            self._reset_chat_history_()
            answer = self._chat_(question)
            doc.add_heading(question, level=2)
            doc.add_heading(f"{llm_name} + {retriever_name} + {filter_name}", level=3)
            doc.add_paragraph(answer)
            doc.add_page_break()
        doc.save(file_path)

    def full_benchmark(self):
            total_time = 0
            total_ttft = 0
            questions = all_questions
            iterations = 10

            for question in questions:
                 for i in range(0, iterations):
                    self._reset_chat_history_()
                    print(f"Question: {question}")
                    start_time = time.time()

                    self._chat_(question)
                    completion_time = time.time()

                    time_first_token = completion_time - start_time

                    iteration_time = time.time() - start_time
                    total_time += iteration_time
                    total_ttft += time_first_token
                    
            average_time = total_time / (len(questions) * iterations)
            average_ttft = total_ttft / (len(questions) * iterations) 
            return average_time, average_ttft
    

    def _filter_(self, query):
            model_name = "mixedbread-ai/mxbai-embed-large-v1"
            tokenizer = AutoTokenizer.from_pretrained(model_name)
            model = AutoModel.from_pretrained(model_name)

            def encode(texts):
                tokens = tokenizer(texts, padding=True, truncation=True, return_tensors="pt")
                with torch.no_grad():
                    outputs = model(**tokens)
                embeddings = outputs.last_hidden_state
                attention_mask = tokens["attention_mask"]
                mask = attention_mask.unsqueeze(-1).expand(embeddings.size()).float()
                summed = torch.sum(embeddings * mask, 1)
                counts = torch.clamp(mask.sum(1), min=1e-9)
                return summed / counts

            def normalize(input):
                return torch.nn.functional.normalize(input, p=2, dim=1)

            contexts = self.retriever.retrieve(query, top_k=10)
            if isinstance(contexts, list):
                contexts = [str(chunk) for chunk in contexts]

            weights = [1.0 - (i * 0.1) for i in range(len(contexts))]

            query_embedding = normalize(encode([query]))
            context_embeddings = normalize(encode(contexts))

            weighted_contexts = [weight * context_embedding for weight, context_embedding in zip(weights, context_embeddings)]

            weighted_contexts = torch.stack(weighted_contexts)

            cos_sim = cosine_similarity(query_embedding, weighted_contexts)

            context_scores = []
            for i, ctx in enumerate(contexts):
                context_score = (i, ctx, cos_sim[0][i])
                context_scores.append(context_score)

            context_scores.sort(key=lambda x: x[2], reverse=True)

            context_scores_list = []
            for rank, (i, ctx, score) in enumerate(context_scores, 1):
                if score >= 0.7500:
                    context_scores_list.append(ctx)

            return context_scores_list
                
    def _reset_chat_history_(self):
                self.chat_history = [
                    {"role": "system", "content": "You are an AI assistant that provides answers based on the given context."}
                    ]

    def chat(self):
        print(f"Current LLM: {self.llm}")
        print(f"Current Retriever: {self.retriever.method}")

        print("\nWelcome to the ABBot.\nEnter /help for available commands.")

        while True:
            query = input("Enter your question: ")

            start_time = time.time()

            if query.lower() == "/exit":
                print("Exiting the chat.")
                break

            elif query.lower() == "/reset":
                self._reset_chat_history_()
                print("Chat history reset.")

            elif query.lower() == "/help":

                print(f"""Available commands:
                      - /exit: Exit the chat
                      - /reset: Reset the chat history
                      - /help: Show this help message
                      - /benchmark: Run a benchmark test
                      - /cls or /clear: Clear the output""")

            elif query.lower() == "/benchmark":
                 self._benchmark_()
            
            elif query.lower() == "/clear" or query.lower() == "/cls":
                 os.system('cls' if os.name=='nt' else 'clear')

            else:
                self._chat_(query)
                print("-" * 100)
                print(f"Execution time: {time.time() - start_time:.2f} seconds\n\n")

def main():
    if len(sys.argv) != 4:
        print("Usage: python ABBot.py <retriever_method> <llm_choice> <filter/nofilter>")
        print("Current supported retrievers: tfidf, bm25")
        print("Exiting ABBot, please try again.")
        exit()

    db_path = "token_semantic_chunk_mxbai_large"

    retriever_method, llm_choice, filter = sys.argv[1], sys.argv[2], sys.argv[3]
    retriever = Retriever(db_path=db_path, collection_name="langchain", method=retriever_method)

    if filter.lower() == "filter":
        filter = True
    elif filter.lower() == "nofilter":
        filter = False
    else:
        print("Invalid filter option. Use 'filter' or 'nofilter'.")
        exit() 

    abbot = ABBot(retriever=retriever, llm=llm_choice, filter=filter)
    abbot.chat()

def full_benchmark(llms, retrievers, db_paths, filters):
    for llm in llms:
         for retriever in retrievers:
            retreiver_method = Retriever(db_paths=db_paths, method=retriever, collection_name="langchain")
            for filter in filters:
                abbot = ABBot(retriever=retreiver_method, llm=llm, filter=filter)
                average_time, average_ttft = abbot.full_benchmark()
                with open("benchmark_results.csv", "a", newline="") as csvfile:
                        writer = csv.writer(csvfile)
                        writer.writerow([llm, retriever, filter, average_time, average_ttft])


def write_to_docx(llms, retrievers, db, filters):
    for llm in llms:
        for retriever in retrievers:
                retreiver_method = Retriever(db_path=db, method=retriever, collection_name="langchain")
                for filter in filters:
                    abbot = ABBot(retriever=retreiver_method, llm=llm, filter=filter)
                    print(f"{llm} + {retriever} + {db} + {filter}")
                    abbot.write_to_docx(llm, retriever, filter)


if __name__ == "__main__":
    # Run chatbot as normal
    #main()

    # Change LLM names to local Ollama model names
    llms = [
        "Gemma3:1b",
        "Llama3.2:1b",
        "Llama3.2:3b",
    ]

    retrievers = [
        "bm25",
        "tfidf",
    ]
    

    db_paths = "token_semantic_chunk_mxbai_large"
    

    filters = [
        "filter",
        "nofilter"
    ]

    # Run full benchmark on all questions and configurations
    #full_benchmark(llms, retrievers, db_paths="token_semantic_chunk_mxbai_large", filters)

    # Run each configuration and write to docx
    write_to_docx(llms, retrievers, db_paths, filters) # IMPORTANT: the _chat_ function needs to return (not print) the answer to be written to docx.

    